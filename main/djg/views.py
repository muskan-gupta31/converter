import cv2
import os
import uuid
import numpy as np
from django.views import View
from django.shortcuts import render
from django.http import JsonResponse
from django.conf import settings
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.utils.decorators import method_decorator
from django.views.decorators.csrf import csrf_exempt
import pandas as pd
from io import BytesIO
from django.shortcuts import render
from django.http import JsonResponse, HttpResponse
from django.views import View
from django.conf import settings
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.utils.decorators import method_decorator
from django.views.decorators.csrf import csrf_exempt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.shared import Inches
import PyPDF2
import csv
import json
import tempfile


PASSPORT_SIZE = (295, 413)  # (width, height) in pixels approx. passport size at 300dpi
GAP = 5  # gap between photos in pixels


@method_decorator(csrf_exempt, name='dispatch')  # Remove in production for security
class PassportSheetView(View):
    def get(self, request):
        return render(request, "home.html")

    def post(self, request):
        file = request.FILES.get("photo")
        try:
            num_copies = int(request.POST.get("copies", "1"))
            if num_copies < 1 or num_copies > 30:
                return JsonResponse({"error": "Choose 1-30 copies."}, status=400)
        except Exception:
            return JsonResponse({"error": "Invalid copy count."}, status=400)

        if not file:
            return JsonResponse({"error": "No photo uploaded"}, status=400)

        # Save uploaded photo temporarily
        temp_name = f"{uuid.uuid4().hex}.jpg"
        temp_path = os.path.join(settings.MEDIA_ROOT, temp_name)
        default_storage.save(temp_name, ContentFile(file.read()))

        # Read photo and validate
        img = cv2.imread(temp_path)
        if img is None:
            default_storage.delete(temp_name)
            return JsonResponse({"error": "Invalid image file."}, status=400)

        # Center crop to square and resize passport size
        h, w = img.shape[:2]
        size = min(h, w)
        center_y, center_x = h // 2, w // 2
        cropped = img[center_y - size // 2:center_y + size // 2, center_x - size // 2:center_x + size // 2]
        passport_img = cv2.resize(cropped, PASSPORT_SIZE)

        # Prepare A4 sheet at 300 dpi: 2480x3508 pixels (WxH)
        SHEET_SIZE = (2480, 3508)
        sheet = np.ones((SHEET_SIZE[1], SHEET_SIZE[0], 3), dtype=np.uint8) * 255

        # Calculate how many can fit horizontally and vertically (with gap)
        cols = SHEET_SIZE[0] // (PASSPORT_SIZE[0] + GAP)
        rows = SHEET_SIZE[1] // (PASSPORT_SIZE[1] + GAP)
        max_photos = cols * rows
        to_place = min(num_copies, max_photos)

        c = 0
        for r in range(rows):
            for col in range(cols):
                if c >= to_place:
                    break
                y = r * (PASSPORT_SIZE[1] + GAP)
                x = col * (PASSPORT_SIZE[0] + GAP)
                sheet[y:y + PASSPORT_SIZE[1], x:x + PASSPORT_SIZE[0]] = passport_img
                c += 1
            if c >= to_place:
                break

        # Save final sheet image
        final_name = f"sheet_{uuid.uuid4().hex}.jpg"
        final_path = os.path.join(settings.MEDIA_ROOT, final_name)
        cv2.imwrite(final_path, sheet, [int(cv2.IMWRITE_JPEG_QUALITY), 95])

        # Clean up temp uploaded file
        default_storage.delete(temp_name)

        url = settings.MEDIA_URL + final_name
        return JsonResponse({"sheet_url": url})





class ConverterView(View):
    """
    Universal file converter supporting multiple formats:
    - PDF ↔ CSV, Excel, Word, TXT
    - CSV ↔ PDF, Excel, Word, TXT
    - Excel ↔ PDF, CSV, Word, TXT
    - Word ↔ PDF, CSV, Excel, TXT
    """
    
    SUPPORTED_FORMATS = {
        'pdf': ['.pdf'],
        'csv': ['.csv'],
        'excel': ['.xlsx', '.xls'],
        'word': ['.docx', '.doc'],
        'txt': ['.txt']
    }
    
    def get(self, request):
        """Render the converter form"""
        return render(request, 'converter.html')
    
    @method_decorator(csrf_exempt)
    def dispatch(self, *args, **kwargs):
        return super().dispatch(*args, **kwargs)
    
    def post(self, request):
        """Handle file conversion requests"""
        try:
            if 'file' not in request.FILES:
                return JsonResponse({'error': 'No file provided'}, status=400)
            
            uploaded_file = request.FILES['file']
            target_format = request.POST.get('target_format', '').lower()
            
            if not target_format:
                return JsonResponse({'error': 'Target format not specified'}, status=400)
            
            # Detect source format
            source_format = self._detect_format(uploaded_file.name)
            if not source_format:
                return JsonResponse({'error': 'Unsupported source file format'}, status=400)
            
            if target_format not in self.SUPPORTED_FORMATS:
                return JsonResponse({'error': f'Unsupported target format: {target_format}'}, status=400)
            
            if source_format == target_format:
                return JsonResponse({'error': 'Source and target formats are the same'}, status=400)
            
            # Save uploaded file temporarily
            temp_path = self._save_temp_file(uploaded_file)
            
            try:
                # Convert file
                converted_file_path = self._convert_file(temp_path, source_format, target_format)
                
                # Serve converted file
                response = self._serve_converted_file(converted_file_path, target_format)
                return response
                
            finally:
                # Cleanup
                self._cleanup_files([temp_path])
                
        except Exception as e:
            return JsonResponse({'error': f'Conversion failed: {str(e)}'}, status=500)
    
    def _detect_format(self, filename):
        """Detect file format from extension"""
        ext = os.path.splitext(filename.lower())[1]
        for format_name, extensions in self.SUPPORTED_FORMATS.items():
            if ext in extensions:
                return format_name
        return None
    
    def _save_temp_file(self, uploaded_file):
        """Save uploaded file to temporary location"""
        temp_dir = os.path.join(settings.MEDIA_ROOT, 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        
        temp_path = os.path.join(temp_dir, uploaded_file.name)
        with open(temp_path, 'wb') as f:
            for chunk in uploaded_file.chunks():
                f.write(chunk)
        return temp_path
    
    def _convert_file(self, source_path, source_format, target_format):
        """Main conversion logic"""
        conversion_method = f"_convert_{source_format}_to_{target_format}"
        if hasattr(self, conversion_method):
            return getattr(self, conversion_method)(source_path)
        else:
            raise NotImplementedError(f"Conversion from {source_format} to {target_format} not implemented")
    
    # PDF Conversions
    def _convert_pdf_to_csv(self, pdf_path):
        """Extract text from PDF and save as CSV"""
        text_data = self._extract_pdf_text(pdf_path)
        csv_path = pdf_path.replace('.pdf', '_converted.csv')
        
        # Convert text to CSV format (basic implementation)
        lines = text_data.split('\n')
        with open(csv_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['Content'])
            for line in lines:
                if line.strip():
                    writer.writerow([line.strip()])
        return csv_path
    
    def _convert_pdf_to_excel(self, pdf_path):
        """Extract text from PDF and save as Excel"""
        text_data = self._extract_pdf_text(pdf_path)
        excel_path = pdf_path.replace('.pdf', '_converted.xlsx')
        
        lines = [line.strip() for line in text_data.split('\n') if line.strip()]
        df = pd.DataFrame({'Content': lines})
        df.to_excel(excel_path, index=False)
        return excel_path
    
    def _convert_pdf_to_word(self, pdf_path):
        """Extract text from PDF and save as Word document"""
        text_data = self._extract_pdf_text(pdf_path)
        word_path = pdf_path.replace('.pdf', '_converted.docx')
        
        doc = Document()
        doc.add_heading('Converted from PDF', 0)
        doc.add_paragraph(text_data)
        doc.save(word_path)
        return word_path
    
    def _convert_pdf_to_txt(self, pdf_path):
        """Extract text from PDF and save as TXT"""
        text_data = self._extract_pdf_text(pdf_path)
        txt_path = pdf_path.replace('.pdf', '_converted.txt')
        
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(text_data)
        return txt_path
    
    # CSV Conversions
    def _convert_csv_to_pdf(self, csv_path):
        """Convert CSV to PDF table"""
        pdf_path = csv_path.replace('.csv', '_converted.pdf')
        
        df = pd.read_csv(csv_path)
        
        doc = SimpleDocTemplate(pdf_path, pagesize=A4)
        elements = []
        
        # Add title
        styles = getSampleStyleSheet()
        title = Paragraph("CSV Data", styles['Title'])
        elements.append(title)
        
        # Convert DataFrame to table
        data = [df.columns.tolist()] + df.values.tolist()
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(table)
        
        doc.build(elements)
        return pdf_path
    
    def _convert_csv_to_excel(self, csv_path):
        """Convert CSV to Excel"""
        excel_path = csv_path.replace('.csv', '_converted.xlsx')
        df = pd.read_csv(csv_path)
        df.to_excel(excel_path, index=False)
        return excel_path
    
    def _convert_csv_to_word(self, csv_path):
        """Convert CSV to Word document with table"""
        word_path = csv_path.replace('.csv', '_converted.docx')
        df = pd.read_csv(csv_path)
        
        doc = Document()
        doc.add_heading('CSV Data', 0)
        
        # Create table
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        # Add headers
        for i, column in enumerate(df.columns):
            table.cell(0, i).text = str(column)
        
        # Add data
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)
        
        doc.save(word_path)
        return word_path
    
    def _convert_csv_to_txt(self, csv_path):
        """Convert CSV to TXT"""
        txt_path = csv_path.replace('.csv', '_converted.txt')
        df = pd.read_csv(csv_path)
        
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(df.to_string(index=False))
        return txt_path
    
    # Excel Conversions
    def _convert_excel_to_pdf(self, excel_path):
        """Convert Excel to PDF"""
        pdf_path = excel_path.replace('.xlsx', '_converted.pdf').replace('.xls', '_converted.pdf')
        df = pd.read_excel(excel_path)
        
        doc = SimpleDocTemplate(pdf_path, pagesize=A4)
        elements = []
        
        styles = getSampleStyleSheet()
        title = Paragraph("Excel Data", styles['Title'])
        elements.append(title)
        
        data = [df.columns.tolist()] + df.values.tolist()
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(table)
        
        doc.build(elements)
        return pdf_path
    
    def _convert_excel_to_csv(self, excel_path):
        """Convert Excel to CSV"""
        csv_path = excel_path.replace('.xlsx', '_converted.csv').replace('.xls', '_converted.csv')
        df = pd.read_excel(excel_path)
        df.to_csv(csv_path, index=False)
        return csv_path
    
    def _convert_excel_to_word(self, excel_path):
        """Convert Excel to Word document"""
        word_path = excel_path.replace('.xlsx', '_converted.docx').replace('.xls', '_converted.docx')
        df = pd.read_excel(excel_path)
        
        doc = Document()
        doc.add_heading('Excel Data', 0)
        
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        for i, column in enumerate(df.columns):
            table.cell(0, i).text = str(column)
        
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)
        
        doc.save(word_path)
        return word_path
    
    def _convert_excel_to_txt(self, excel_path):
        """Convert Excel to TXT"""
        txt_path = excel_path.replace('.xlsx', '_converted.txt').replace('.xls', '_converted.txt')
        df = pd.read_excel(excel_path)
        
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(df.to_string(index=False))
        return txt_path
    
    # Word Conversions
    def _convert_word_to_pdf(self, word_path):
        """Convert Word to PDF (basic text extraction)"""
        pdf_path = word_path.replace('.docx', '_converted.pdf').replace('.doc', '_converted.pdf')
        
        # Extract text from Word document
        doc = Document(word_path)
        text_content = []
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        
        # Create PDF
        doc_pdf = SimpleDocTemplate(pdf_path, pagesize=A4)
        styles = getSampleStyleSheet()
        elements = []
        
        elements.append(Paragraph("Converted from Word Document", styles['Title']))
        for text in text_content:
            if text.strip():
                elements.append(Paragraph(text, styles['Normal']))
        
        doc_pdf.build(elements)
        return pdf_path
    
    def _convert_word_to_csv(self, word_path):
        """Convert Word to CSV (extract text as rows)"""
        csv_path = word_path.replace('.docx', '_converted.csv').replace('.doc', '_converted.csv')
        
        doc = Document(word_path)
        text_lines = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_lines.append(paragraph.text.strip())
        
        with open(csv_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['Content'])
            for line in text_lines:
                writer.writerow([line])
        
        return csv_path
    
    def _convert_word_to_excel(self, word_path):
        """Convert Word to Excel"""
        excel_path = word_path.replace('.docx', '_converted.xlsx').replace('.doc', '_converted.xlsx')
        
        doc = Document(word_path)
        text_lines = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_lines.append(paragraph.text.strip())
        
        df = pd.DataFrame({'Content': text_lines})
        df.to_excel(excel_path, index=False)
        return excel_path
    
    def _convert_word_to_txt(self, word_path):
        """Convert Word to TXT"""
        txt_path = word_path.replace('.docx', '_converted.txt').replace('.doc', '_converted.txt')
        
        doc = Document(word_path)
        with open(txt_path, 'w', encoding='utf-8') as f:
            for paragraph in doc.paragraphs:
                f.write(paragraph.text + '\n')
        
        return txt_path
    
    # TXT Conversions
    def _convert_txt_to_pdf(self, txt_path):
        """Convert TXT to PDF"""
        pdf_path = txt_path.replace('.txt', '_converted.pdf')
        
        with open(txt_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        doc = SimpleDocTemplate(pdf_path, pagesize=A4)
        styles = getSampleStyleSheet()
        elements = []
        
        elements.append(Paragraph("Text Document", styles['Title']))
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                elements.append(Paragraph(para.replace('\n', ' '), styles['Normal']))
        
        doc.build(elements)
        return pdf_path
    
    def _convert_txt_to_csv(self, txt_path):
        """Convert TXT to CSV"""
        csv_path = txt_path.replace('.txt', '_converted.csv')
        
        with open(txt_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        with open(csv_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['Content'])
            for line in lines:
                if line.strip():
                    writer.writerow([line.strip()])
        
        return csv_path
    
    def _convert_txt_to_excel(self, txt_path):
        """Convert TXT to Excel"""
        excel_path = txt_path.replace('.txt', '_converted.xlsx')
        
        with open(txt_path, 'r', encoding='utf-8') as f:
            lines = [line.strip() for line in f.readlines() if line.strip()]
        
        df = pd.DataFrame({'Content': lines})
        df.to_excel(excel_path, index=False)
        return excel_path
    
    def _convert_txt_to_word(self, txt_path):
        """Convert TXT to Word"""
        word_path = txt_path.replace('.txt', '_converted.docx')
        
        with open(txt_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        doc = Document()
        doc.add_heading('Text Document', 0)
        doc.add_paragraph(content)
        doc.save(word_path)
        return word_path
    
    # Helper methods
    def _extract_pdf_text(self, pdf_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise Exception(f"Error extracting PDF text: {str(e)}")
        return text
    
    def _serve_converted_file(self, file_path, target_format):
        """Serve the converted file as download"""
        filename = os.path.basename(file_path)
        
        content_types = {
            'pdf': 'application/pdf',
            'csv': 'text/csv',
            'excel': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'word': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'txt': 'text/plain'
        }
        
        with open(file_path, 'rb') as f:
            response = HttpResponse(f.read(), content_type=content_types.get(target_format, 'application/octet-stream'))
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
    
    def _cleanup_files(self, file_paths):
        """Clean up temporary files"""
        for path in file_paths:
            try:
                if os.path.exists(path):
                    os.remove(path)
            except:
                pass